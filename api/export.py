import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.text.run import Font
from convert_to_pdf import convert

def export_pdf(user):
    # get an unused filename in current cache
    filename = ""

    doc = docx.Document()
    doc.sections[0].top_margin = Inches(1)
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].bottom_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)

    styles = doc.styles
    for i, s in enumerate(styles):
        print(i, s.name)

    # Generate header
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False 
    table.allow_autofit = False

    WD_ALIGN_PARAGRAPH.CENTER
    WD_ALIGN_VERTICAL.CENTER

    table.style = styles['Colorful Grid Accent 4']

    logo = table.rows[0].cells[0].paragraphs[0].add_run().add_picture('assets/Sample_logo.png')
    logo.width = Inches(1.4)
    logo.height = Inches(0.7)
    table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = Inches(1.6)

    table.rows[0].cells[1].text = user['name']
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(24)
    table.columns[1].width = Inches(3.3)
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    qr = table.rows[0].cells[2].paragraphs[0].add_run().add_picture('assets/Sample_qr.png')
    qr.width = Inches(0.7)
    qr.height = Inches(0.7)
    table.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.columns[2].width = Inches(1.6)

    table.rows[0].height = Inches(0.8)

    # Generate disabilities

    p = doc.add_paragraph('This is a paragraph2')
    p = doc.add_paragraph('This is a paragraph2')

    # TODO remove in production
    doc.save('demo.docx')
    convert(doc)


def export_by_type(user: str, ftype: str):
    if ftype == 'pdf':
        return export_pdf(user)
    if ftype == 'card':
        return None
    if ftype == 'docx':
        return None
    if ftype == 'png':
        return None
    if ftype == 'qr':
        return None
    

export_pdf({'name': 'John Smith'})
